"""
utils/file_handler.py - Extract plain text from uploaded files (PDF, DOCX, TXT)

Uses a multi-strategy approach for PDFs:
  1. pdfplumber  — best for structured / multi-column resumes
  2. PyPDF2      — fallback for simpler PDFs
  3. Gemini multimodal — fallback for scanned / image-heavy PDFs
"""

import io
import os
import re
import base64
import PyPDF2
import pdfplumber
import docx
import google.generativeai as genai
from werkzeug.datastructures import FileStorage
from dotenv import load_dotenv

load_dotenv()

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "doc"}
MAX_FILE_SIZE_MB = 10
MIN_USEFUL_TEXT_LENGTH = 50  # below this, consider extraction failed


def _allowed_file(filename: str) -> bool:
    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS
    )


# ── PDF Extraction Strategies ───────────────────────────────────────────────────

def _extract_pdf_pdfplumber(file_bytes: bytes) -> str:
    """Strategy 1: pdfplumber — excellent for complex layouts & tables."""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text_parts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts).strip()
    except Exception as e:
        print(f"  ℹ️  pdfplumber failed: {e}")
        return ""


def _extract_pdf_pypdf2(file_bytes: bytes) -> str:
    """Strategy 2: PyPDF2 — simpler extraction, good fallback."""
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text_parts.append(extracted)
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"  ℹ️  PyPDF2 failed: {e}")
        return ""


def _extract_pdf_gemini(file_bytes: bytes) -> str:
    """
    Strategy 3: Gemini multimodal — send PDF as inline data.
    Handles scanned PDFs and image-heavy documents via OCR-like extraction.
    """
    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            return ""

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.5-flash")

        # Encode PDF as base64 for inline data
        b64_pdf = base64.b64encode(file_bytes).decode("utf-8")

        prompt = """You are a document text extractor. Extract ALL the text content from 
this PDF document. Preserve the structure as much as possible — sections, bullet points, 
dates, job titles, skills, education, etc. Return ONLY the extracted plain text, 
nothing else. Do not add any commentary or formatting markers."""

        response = model.generate_content([
            prompt,
            {
                "mime_type": "application/pdf",
                "data": b64_pdf,
            }
        ])

        if response and response.text:
            return response.text.strip()
        return ""
    except Exception as e:
        print(f"  ℹ️  Gemini PDF extraction failed: {e}")
        return ""


def _clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ""
    # Collapse excessive whitespace but preserve paragraph breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    # Remove null bytes and other control chars
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    return text.strip()


def _extract_from_pdf(file_bytes: bytes) -> str:
    """
    Multi-strategy PDF extraction with quality detection.
    Tries pdfplumber → PyPDF2 → Gemini multimodal, picks the best result.
    """
    print("  📄 Extracting text from PDF...")

    # Strategy 1: pdfplumber (best for most resumes)
    text_plumber = _clean_extracted_text(_extract_pdf_pdfplumber(file_bytes))
    if len(text_plumber) >= MIN_USEFUL_TEXT_LENGTH:
        print(f"  ✅ pdfplumber extracted {len(text_plumber)} chars")
        return text_plumber

    # Strategy 2: PyPDF2
    text_pypdf2 = _clean_extracted_text(_extract_pdf_pypdf2(file_bytes))
    if len(text_pypdf2) >= MIN_USEFUL_TEXT_LENGTH:
        print(f"  ✅ PyPDF2 extracted {len(text_pypdf2)} chars")
        return text_pypdf2

    # Strategy 3: Gemini multimodal (for scanned/image PDFs)
    print("  🔄 Text extractors yielded little text — trying Gemini multimodal...")
    text_gemini = _clean_extracted_text(_extract_pdf_gemini(file_bytes))
    if len(text_gemini) >= MIN_USEFUL_TEXT_LENGTH:
        print(f"  ✅ Gemini multimodal extracted {len(text_gemini)} chars")
        return text_gemini

    # Return whatever we have (even if short)
    best = max([text_plumber, text_pypdf2, text_gemini], key=len)
    if best:
        print(f"  ⚠️  Best extraction only {len(best)} chars — quality may be low")
        return best

    return ""


def _extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes — includes tables and headers."""
    doc = docx.Document(io.BytesIO(file_bytes))

    parts = []

    # Extract paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # Also extract text from tables (resumes often use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def _extract_from_txt(file_bytes: bytes) -> str:
    """Decode plain text bytes."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace").strip()


def extract_text_from_file(file: FileStorage) -> str:
    """
    Extract plain text from an uploaded Werkzeug FileStorage object.

    Supports: PDF, DOCX, TXT
    Raises:   ValueError for unsupported types or empty results.
    """
    filename = file.filename or ""

    if not _allowed_file(filename):
        raise ValueError(
            f"Unsupported file type '{filename}'. "
            f"Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    file_bytes = file.read()

    # Guard against empty files
    if not file_bytes:
        raise ValueError(f"Uploaded file '{filename}' is empty.")

    # Guard against oversized files
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(
            f"File '{filename}' is {size_mb:.1f} MB — "
            f"max allowed is {MAX_FILE_SIZE_MB} MB."
        )

    ext = filename.rsplit(".", 1)[1].lower()

    if ext == "pdf":
        text = _extract_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        text = _extract_from_docx(file_bytes)
    elif ext == "txt":
        text = _extract_from_txt(file_bytes)
    else:
        raise ValueError(f"Unhandled extension: {ext}")

    if not text:
        raise ValueError(
            f"Could not extract any text from '{filename}'. "
            "The file may be scanned/image-only or corrupt."
        )

    return text


def extract_text_from_file_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from raw bytes + filename (for non-Werkzeug contexts).
    """
    if not file_bytes:
        raise ValueError("File bytes are empty.")

    ext = filename.rsplit(".", 1)[1].lower() if "." in filename else ""

    if ext == "pdf":
        return _extract_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return _extract_from_docx(file_bytes)
    elif ext == "txt":
        return _extract_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported extension: {ext}")


def get_text_input(
    request_form: dict,
    request_files: dict,
    text_field: str,
    file_field: str,
) -> str:
    """
    Resolve text input from either a form text field or an uploaded file.

    Priority: uploaded file > plain text field.

    Returns:
        Extracted/provided text string.
    Raises:
        ValueError if neither source is provided or both are empty.
    """
    # Check for uploaded file first
    if file_field in request_files:
        uploaded = request_files[file_field]
        if uploaded and uploaded.filename:
            return extract_text_from_file(uploaded)

    # Fall back to plain text field
    text = request_form.get(text_field, "").strip()
    if text:
        return text

    raise ValueError(
        f"No input provided for '{text_field}'. "
        f"Send either a '{text_field}' text field or a '{file_field}' file."
    )
